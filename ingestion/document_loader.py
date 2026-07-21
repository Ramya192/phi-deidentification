"""
ingestion/document_loader.py

Extracts raw text from an uploaded file (.txt, .pdf, .docx) before it
enters the LangGraph pipeline via ClassificationAgent.

This lives OUTSIDE the LangGraph graph on purpose. GraphState
(graph/state.py) is checkpointed via LangGraph's MemorySaver to support
HumanReviewAgent's interrupt()/resume flow, and keeping that state a
plain JSON-serializable dict (raw_text: str, not binary blobs) keeps
checkpointing simple and keeps every agent testable with plain strings.
So extraction happens once, up front — see api/main.py's
POST /redact/upload — and only the extracted text (plus the original
filename, for the audit trail) crosses into GraphState.

Supported formats: .txt, .pdf, .docx
Not supported: legacy .doc, scanned/image-only PDFs without OCR (see
extract_pdf_with_ocr, which is opt-in and not wired into the default
path since it needs system-level binaries, not just pip packages).

USAGE
-----
    from ingestion.document_loader import load_text_from_bytes

    text = load_text_from_bytes(file_bytes, filename="discharge_summary.pdf")
"""
from __future__ import annotations

import io
import os

SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


class UnsupportedFileTypeError(ValueError):
    """Raised when the file extension isn't one we know how to extract text from."""


class EmptyDocumentError(ValueError):
    """Raised when extraction succeeded but produced no usable text (e.g. a
    blank file, or a scanned PDF with no embedded text layer)."""


class OCRNotAvailableError(RuntimeError):
    """Raised when extract_pdf_with_ocr() is called but its dependencies
    (pytesseract, pdf2image, and the tesseract/poppler system binaries)
    aren't installed."""


def load_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """Extracts text from raw file bytes based on the filename's extension.

    Raises UnsupportedFileTypeError for unrecognized extensions, and
    EmptyDocumentError if extraction produced no usable text.
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".txt":
        text = _extract_txt(file_bytes)
    elif ext == ".pdf":
        text = _extract_pdf(file_bytes)
    elif ext == ".docx":
        text = _extract_docx(file_bytes)
    else:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{ext}' for '{filename}'. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}. "
            f"(Legacy .doc is not supported — convert to .docx or .pdf first.)"
        )

    text = text.strip()
    if not text:
        raise EmptyDocumentError(
            f"No extractable text found in '{filename}'. If this is a scanned "
            f"PDF (image-only, no text layer), OCR is required — see "
            f"extract_pdf_with_ocr() in this module."
        )
    return text


def load_text_from_path(path: str) -> str:
    """Convenience wrapper for CLI/test use — reads the file off disk and
    delegates to load_text_from_bytes()."""
    with open(path, "rb") as f:
        file_bytes = f.read()
    return load_text_from_bytes(file_bytes, filename=os.path.basename(path))


def _extract_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode .txt file — tried utf-8 and latin-1.")


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        import pdfplumber
    except ImportError as exc:
        raise ImportError(
            "PDF support requires: pip install pdfplumber"
        ) from exc

    pages_text: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            pages_text.append(page.extract_text() or "")
    return "\n\n".join(pages_text)


def _extract_docx(file_bytes: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise ImportError(
            "DOCX support requires: pip install python-docx"
        ) from exc

    document = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def extract_pdf_with_ocr(file_bytes: bytes, dpi: int = 300) -> str:
    """Fallback for scanned/image-only PDFs with no embedded text layer.

    Opt-in only — not called automatically by load_text_from_bytes(),
    because it needs the tesseract-ocr and poppler system binaries
    installed (not just pip packages), is much slower than text
    extraction, and shouldn't silently fire on every PDF. Call this
    explicitly when _extract_pdf() raises EmptyDocumentError and you know
    the source is a scan.

    Install (in addition to `pip install pytesseract pdf2image`):
      - macOS:   brew install tesseract poppler
      - Ubuntu:  apt-get install tesseract-ocr poppler-utils
      - Windows: install Tesseract-OCR and Poppler binaries separately
                 and add both to PATH.
    """
    try:
        import pytesseract
        from pdf2image import convert_from_bytes
    except ImportError as exc:
        raise OCRNotAvailableError(
            "OCR requires: pip install pytesseract pdf2image, plus the "
            "tesseract-ocr and poppler system packages (not pip-installable "
            "— see this function's docstring for platform-specific installs)."
        ) from exc

    images = convert_from_bytes(file_bytes, dpi=dpi)
    pages_text = [pytesseract.image_to_string(img) for img in images]
    text = "\n\n".join(pages_text).strip()
    if not text:
        raise EmptyDocumentError("OCR ran but found no text in this PDF.")
    return text
